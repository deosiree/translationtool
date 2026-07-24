#!/usr/bin/env python3
"""Export styles/*.md into formats/raw/{md,txt,docx,pdf} multi-format families for cleaning drills."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
STYLES = ROOT / "data" / "rag-corpus" / "styles"
RAW = ROOT / "data" / "rag-corpus" / "formats" / "raw"
MANIFEST = ROOT / "data" / "rag-corpus" / "MANIFEST.yaml"

# Keep samples short for cleaning drills (guide: 1–3 pages). Cap han chars.
MAX_HAN = 3500
HAN_RE = re.compile(r"[\u4e00-\u9fff]")

# (styles relative path, content_family, module, stem for filenames)
EXPORTS = [
    ("faq/工作台常见问题.md", "workbench-faq", "workbench", "工作台常见问题"),
    ("faq/词条管理常见问题.md", "entry-faq", "entry", "词条管理常见问题"),
    ("faq/术语库常见问题.md", "glossary-faq", "glossary", "术语库常见问题"),
    ("faq/配置管理常见问题.md", "configure-faq", "configure", "配置管理常见问题"),
    ("faq/翻译校验常见问题.md", "check-faq", "check", "翻译校验常见问题"),
    ("faq/文件管理常见问题.md", "file-faq", "file", "文件管理常见问题"),
    ("faq/术语学习常见问题.md", "termlearn-faq", "terminologyAgent", "术语学习常见问题"),
    ("faq/悬浮工具箱常见问题.md", "toolbox-faq", "toolbox", "悬浮工具箱常见问题"),
    ("faq/智能助手常见问题.md", "assistant-faq", "assistant", "智能助手常见问题"),
    ("troubleshooting/工作台排障.md", "workbench-trouble", "workbench", "工作台排障"),
    ("troubleshooting/配置管理排障.md", "configure-trouble", "configure", "配置管理排障"),
    ("sop/workbench/词条导入.md", "sop-wb-import", "workbench", "词条导入SOP"),
    ("sop/workbench/词条翻译.md", "sop-wb-translate", "workbench", "词条翻译SOP"),
    ("sop/configure/任务创建与下发.md", "sop-cfg-task", "configure", "任务创建与下发SOP"),
    ("user-manual/工作台使用说明.md", "workbench-manual", "workbench", "工作台使用说明"),
    ("user-manual/配置管理使用说明.md", "configure-manual", "configure", "配置管理使用说明"),
    ("architecture/翻译工具全系统架构.md", "system-architecture", "platform", "翻译工具全系统架构"),
    ("ops-notes/本地启动与端口备忘.md", "ops-ports", "ops", "本地启动与端口备忘"),
    ("scenarios/工作台场景01.md", "workbench-scenario-01", "workbench", "工作台场景01"),
    ("scenarios/术语学习场景01.md", "termlearn-scenario-01", "terminologyAgent", "术语学习场景01"),
]


def han_count(s: str) -> int:
    return len(HAN_RE.findall(s))


def clip_md(text: str, max_han: int = MAX_HAN) -> str:
    if han_count(text) <= max_han:
        return text
    # clip by paragraphs
    out = []
    n = 0
    for para in text.split("\n"):
        c = han_count(para)
        if n + c > max_han and n > max_han * 0.7:
            out.append("\n\n> （导出截断：完整版见 styles/ 同源文档）\n")
            break
        out.append(para)
        n += c
    return "\n".join(out)


def find_cn_font() -> str | None:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttf"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return None


def write_txt(path: Path, text: str) -> None:
    # strip markdown-ish for cleaning challenge
    plain = re.sub(r"^#+\s*", "", text, flags=re.M)
    plain = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", plain)
    plain = plain.replace("**", "").replace("`", "")
    path.write_text(plain, encoding="utf-8")


def write_docx(path: Path, title: str, text: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.splitlines():
        if line.startswith("#"):
            level = min(line.count("#", 0, 6), 3)
            doc.add_heading(line.lstrip("# ").strip() or title, level=level)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    doc.save(path)


def write_pdf(path: Path, title: str, text: str, font_path: str | None) -> None:
    if not font_path:
        # fallback: ascii-only note + utf-8 sidecar already in md/txt
        path.write_bytes(b"%PDF-1.4\n% fallback empty\n")
        raise RuntimeError("no Chinese font for PDF")

    font_name = "CNFont"
    try:
        pdfmetrics.registerFont(TTFont(font_name, font_path, subfontIndex=0))
    except Exception:
        pdfmetrics.registerFont(TTFont(font_name, font_path))

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CNBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )
    h1 = ParagraphStyle(
        "CNH1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        leading=18,
    )

    def esc(s: str) -> str:
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    story = [Paragraph(esc(title), h1), Spacer(1, 12)]
    for line in text.splitlines():
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        if line.startswith("#"):
            story.append(Paragraph(esc(line.lstrip("# ").strip()), h1))
        else:
            story.append(Paragraph(esc(line), body))
            story.append(Spacer(1, 4))

    SimpleDocTemplate(str(path), pagesize=A4).build(story)


def main() -> None:
    font = find_cn_font()
    print("font:", font)
    for sub in ("md", "txt", "docx", "pdf"):
        (RAW / sub).mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict] = []
    ok = 0
    for rel, family, module, stem in EXPORTS:
        src = STYLES / rel
        if not src.is_file():
            print("SKIP missing", rel)
            continue
        text = clip_md(src.read_text(encoding="utf-8"))
        # md mirror
        md_path = RAW / "md" / f"{stem}.md"
        md_path.write_text(text, encoding="utf-8")
        txt_path = RAW / "txt" / f"{stem}.txt"
        write_txt(txt_path, text)
        docx_path = RAW / "docx" / f"{stem}.docx"
        write_docx(docx_path, stem, text)
        pdf_path = RAW / "pdf" / f"{stem}.pdf"
        try:
            write_pdf(pdf_path, stem, text, font)
        except Exception as e:
            print("PDF fail", stem, e)
            continue

        for fmt, path in (
            ("md", md_path),
            ("txt", txt_path),
            ("docx", docx_path),
            ("pdf", pdf_path),
        ):
            challenges = {
                "md": ["markdown_headers", "screenshot_placeholders"],
                "txt": ["lost_structure", "plain_paragraphs"],
                "docx": ["docx_styles", "heading_levels"],
                "pdf": ["pdf_extract", "cjk_font"],
            }[fmt]
            manifest_rows.append(
                {
                    "id": f"fmt-{family}-{fmt}",
                    "kind": "format",
                    "path": path.relative_to(ROOT / "data" / "rag-corpus").as_posix(),
                    "style": "raw-format",
                    "module": module,
                    "format": fmt,
                    "content_family": family,
                    "source_skill": "export",
                    "cleaning_challenges": challenges,
                    "audience": "dev",
                    "status": "ready",
                    "source_style_path": f"styles/{rel}",
                }
            )
        ok += 1
        print("OK", stem, "han", han_count(text))

    # merge manifest: drop old format entries for these families then append
    import yaml

    data = yaml.safe_load(MANIFEST.read_text(encoding="utf-8")) or {}
    entries = data.get("entries") or []
    families = {r["content_family"] for r in manifest_rows}
    kept = [
        e
        for e in entries
        if not (
            e.get("kind") == "format"
            and e.get("content_family") in families
            and str(e.get("path", "")).startswith("formats/raw/")
        )
    ]
    # also drop obsolete assistant-faq format paths that we recreate
    data["entries"] = kept + manifest_rows
    MANIFEST.write_text(yaml.dump(data, allow_unicode=True, sort_keys=False), encoding="utf-8")
    print(f"exported families={ok} format_entries={len(manifest_rows)}")


if __name__ == "__main__":
    main()
